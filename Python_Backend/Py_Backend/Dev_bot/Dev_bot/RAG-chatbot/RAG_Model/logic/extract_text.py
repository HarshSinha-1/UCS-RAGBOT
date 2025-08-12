from PyPDF2 import PdfReader
from docx import Document
from bs4 import BeautifulSoup
import pdfplumber
import os
from PIL import Image
import pytesseract
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
import io
import csv
import json
from pathlib import Path
from dotenv import load_dotenv
load_dotenv()
import fitz  # PyMuPDF
import sys
os.environ["TOKENIZERS_PARALLELISM"] = "false"

import warnings
warnings.filterwarnings("ignore", category=UserWarning)



# def extract_text_from_pdf(file_path):
#     extracted = []

#     # --- Text and Table Extraction with pdfplumber ---
#     try:
#         with pdfplumber.open(file_path) as pdf:
#             for i, page in enumerate(pdf.pages):
#                 # Plain text
#                 page_text = page.extract_text()
#                 if page_text:
#                     extracted.append(page_text)

#                 # Tables as Markdown
#                 tables = page.extract_tables()
#                 for t_idx, table in enumerate(tables):
#                     if table and any(any(cell for cell in row) for row in table):
#                         table_md = []
#                         for row in table:
#                             row = [cell if cell is not None else "" for cell in row]
#                             table_md.append('| ' + ' | '.join(row) + ' |')
#                         extracted.append(f"\n[Extracted Table {i+1}.{t_idx+1}]\n" + "\n".join(table_md))
#     except Exception as e:
#         print(f"Error extracting text/tables from PDF: {e}", file=sys.stderr)

#     # --- Image Extraction and OCR with PyMuPDF ---
#     try:
#         doc = fitz.open(file_path)
#         for page_num in range(len(doc)):
#             page = doc[page_num]
#             image_list = page.get_images(full=True)
#             for img_idx, img in enumerate(image_list):
#                 xref = img[0]
#                 base_image = doc.extract_image(xref)
#                 image_bytes = base_image["image"]
#                 try:
#                     pil_img = Image.open(io.BytesIO(image_bytes))
#                     ocr_text = pytesseract.image_to_string(pil_img)
#                     if ocr_text.strip():
#                         extracted.append(f"[OCR from image on page {page_num+1}]\n{ocr_text.strip()}")
#                 except Exception as ocr_e:
#                     extracted.append(f"[Image on page {page_num+1} found, but OCR failed: {ocr_e}]")
#     except Exception as e:
#         print(f"Error extracting images from PDF: {e}", file=sys.stderr)

#     return "\n\n".join(extracted)

# def extract_text_from_pdf(file_path):
#     extracted = []

#     # Step 1: Text + Table extraction using pdfplumber
#     try:
#         with pdfplumber.open(file_path) as pdf:
#             for i, page in enumerate(pdf.pages):
#                 page_chunks = []

#                 # Extract text
#                 page_text = page.extract_text()
#                 if page_text:
#                     page_chunks.append(page_text)

#                 # Extract tables
#                 tables = page.extract_tables()
#                 for t_idx, table in enumerate(tables):
#                     if table and any(any(cell for cell in row) for row in table):
#                         table_md = []
#                         for row in table:
#                             row = [cell if cell is not None else "" for cell in row]
#                             table_md.append('| ' + ' | '.join(row) + ' |')
#                         table_text = f"\n[Extracted Table {i+1}.{t_idx+1}]\n" + "\n".join(table_md)
#                         page_chunks.append(table_text)

#                 full_text = "\n\n".join(page_chunks).strip()
#                 if full_text:
#                     extracted.append((full_text, i + 1))
#                 else:
#                     extracted.append(("", i + 1))  # Ensure page exists for alignment with OCR
#     except Exception as e:
#         print(f"Error extracting text/tables from PDF: {e}", file=sys.stderr)

#     # Step 2: OCR from images using fitz
#     try:
#         doc = fitz.open(file_path)
#         for page_num in range(len(doc)):
#             page = doc[page_num]
#             image_list = page.get_images(full=True)

#             ocr_chunks = []

#             for img_idx, img in enumerate(image_list):
#                 xref = img[0]
#                 base_image = doc.extract_image(xref)
#                 image_bytes = base_image["image"]

#                 try:
#                     pil_img = Image.open(io.BytesIO(image_bytes))
#                     ocr_text = pytesseract.image_to_string(pil_img)
#                     if ocr_text.strip():
#                         ocr_chunks.append(f"[OCR from image on page {page_num+1}]\n{ocr_text.strip()}")
#                 except Exception as ocr_e:
#                     ocr_chunks.append(f"[Image on page {page_num+1} found, but OCR failed: {ocr_e}]")

#             if ocr_chunks:
#                 if page_num < len(extracted):
#                     extracted[page_num] = (
#                         extracted[page_num][0] + "\n\n" + "\n\n".join(ocr_chunks),
#                         page_num + 1
#                     )
#                 else:
#                     extracted.append(("\n\n".join(ocr_chunks), page_num + 1))
#     except Exception as e:
#         print(f"Error extracting images from PDF: {e}", file=sys.stderr)

#     return extracted  # List of tuples: (text, page_number)


# # --- Main document loader ---
# def load_documents(docs_dir):
    
#     docs_dir = Path(docs_dir)  # ✅ Convert str to Path if needed
#     extractors = {
#         '.pdf': extract_text_from_pdf,
#         '.docx': extract_text_from_docx,
#         '.txt': extract_text_from_txt,
#         '.md': extract_text_from_md,
#         '.html': extract_text_from_html,
#         '.htm': extract_text_from_html,
#         '.csv': extract_text_from_csv,
#     }
#     documents = []
#     for file_path in docs_dir.glob('*'):
#         if file_path.suffix.lower() == ".pdf":
#             text = extract_text_from_pdf(str(file_path))
#             #documents.extend([(text, page) for text, page in pagewise])
#             documents.append(text)
#         else:
#             extractor = extractors.get(file_path.suffix.lower())
#             if extractor:
#                 text = extractor(str(file_path))
#                 if text:
#                     documents.append((text, 1))  # Default page_no = 1
#                     print(f"\n--- Extracted from {file_path} ---\n{text[:1000]}...\n")
#                 else:
#                     print(f"Unsupported file type: {file_path.suffix}", file=sys.stderr)
#     return documents  # Always return (text, page_no)



# def extract_text_from_pdf(file_path):
#     extracted = []

#     # --- Text and Table Extraction with pdfplumber ---
#     try:
#         with pdfplumber.open(file_path) as pdf:
#             for i, page in enumerate(pdf.pages):
#                 # Plain text
#                 page_text = page.extract_text()
#                 if page_text:
#                     extracted.append(page_text)

#                 # Tables as Markdown
#                 tables = page.extract_tables()
#                 for t_idx, table in enumerate(tables):
#                     if table and any(any(cell for cell in row) for row in table):
#                         table_md = []
#                         for row in table:
#                             row = [cell if cell is not None else "" for cell in row]
#                             table_md.append('| ' + ' | '.join(row) + ' |')
#                         extracted.append(f"\n[Extracted Table {i+1}.{t_idx+1}]\n" + "\n".join(table_md))
#     except Exception as e:
#         print(f"Error extracting text/tables from PDF: {e}", file=sys.stderr)

#     # --- Image Extraction and OCR with PyMuPDF ---
#     try:
#         doc = fitz.open(file_path)
#         for page_num in range(len(doc)):
#             page = doc[page_num]
#             image_list = page.get_images(full=True)
#             for img_idx, img in enumerate(image_list):
#                 xref = img[0]
#                 base_image = doc.extract_image(xref)
#                 image_bytes = base_image["image"]
#                 try:
#                     pil_img = Image.open(io.BytesIO(image_bytes))
#                     ocr_text = pytesseract.image_to_string(pil_img)
#                     if ocr_text.strip():
#                         extracted.append(f"[OCR from image on page {page_num+1}]\n{ocr_text.strip()}")
#                 except Exception as ocr_e:
#                     extracted.append(f"[Image on page {page_num+1} found, but OCR failed: {ocr_e}]")
#     except Exception as e:
#         print(f"Error extracting images from PDF: {e}", file=sys.stderr)

#     return "\n\n".join(extracted)

# def extract_text_from_pdf(file_path):
#     extracted = []

#     # --- Extract text & tables ---
#     try:
#         print('Text Extraction Started')
#         with pdfplumber.open(file_path) as pdf:
#             for i, page in enumerate(pdf.pages):
#                 page_chunks = []
#                 text = page.extract_text()

#                 if text:
#                     page_chunks.append(text.strip())

#                 tables = page.extract_tables()
#                 for t_idx, table in enumerate(tables):
#                     if table and any(any(cell for cell in row) for row in table):
#                         table_md = []
#                         for row in table:
#                             row = [cell if cell is not None else "" for cell in row]
#                             table_md.append('| ' + ' | '.join(row) + ' |')
#                         table_text = f"\n[Extracted Table {i+1}.{t_idx+1}]\n" + "\n".join(table_md)
#                         page_chunks.append(table_text)

#                 full_text = "\n\n".join(page_chunks).strip()
#                 if full_text:
#                     extracted.append((full_text, i + 1))
#     except Exception as e:
#         print(f"[pdfplumber] Error: {e}", file=sys.stderr)

#     # --- OCR from images ---
#     try:
#         doc = fitz.open(file_path)
#         for page_num in range(len(doc)):
#             page = doc[page_num]
#             image_list = page.get_images(full=True)

#             for img_idx, img in enumerate(image_list):
#                 xref = img[0]
#                 base_image = doc.extract_image(xref)
#                 image_bytes = base_image["image"]

#                 try:
#                     pil_img = Image.open(io.BytesIO(image_bytes)).convert("RGB")
#                     ocr_text = pytesseract.image_to_string(pil_img)
#                     if ocr_text.strip():
#                         extracted.append((f"[OCR Image Text Page {page_num+1}]\n{ocr_text.strip()}", page_num + 1))
#                 except Exception as ocr_e:
#                     print(f"[OCR Failed] Page {page_num+1}: {ocr_e}", file=sys.stderr)
#     except Exception as e:
#         print(f"[fitz] Image extraction error: {e}", file=sys.stderr)

#     print("Extraction ends")

#     return extracted  # list of (text, page_no)

import hashlib
import logging
from typing import List, Tuple


logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def extract_text_from_pdf(
    file_path: Path,
    ocr: bool = True,
    ocr_dpi: int = 300
) -> List[Tuple[str, int]]:
    """
    Extracts text from each page of a PDF.  
    Falls back to OCR when page text is empty.  
    Deduplicates pages by MD5 hash of their text.

    :param file_path: Path to the PDF file.
    :param ocr: whether to run OCR on blank pages.
    :param ocr_dpi: DPI to render images for OCR.
    :return: List of (page_text, page_number).
    """
    extracted: List[Tuple[str, int]] = []
    seen_hashes = set()

    try:
        logger.info("Opening PDF %s", file_path)
        doc = fitz.open(file_path)

        for page_index in range(len(doc)):
            page = doc.load_page(page_index)
            text = page.get_text("text").strip()

            if not text and ocr:
                logger.debug("Page %d empty—running OCR", page_index + 1)
                pix = page.get_pixmap(dpi=ocr_dpi)
                img = Image.open(io.BytesIO(pix.tobytes("png"))).convert("RGB")
                text = pytesseract.image_to_string(img).strip()

            if not text:
                logger.debug("Page %d still empty—skipping", page_index + 1)
                continue

            hash_val = hashlib.md5(text.encode("utf-8")).hexdigest()
            if hash_val in seen_hashes:
                logger.debug("Duplicate page %d skipped", page_index + 1)
                continue

            extracted.append((text, page_index + 1))
            seen_hashes.add(hash_val)

    except Exception as e:
        logger.error("Failed extracting %s: %s", file_path, e, exc_info=True)
        raise

    logger.info("Extraction complete—%d pages", len(extracted))
    return extracted


def extract_text_from_docx(file_path):
    
    doc = Document(file_path)
    extracted = []

    virtual_page_no = 1  # simulate a page count
    content_chunks = []

    # --- Extract paragraphs ---
    for para in doc.paragraphs:
        txt = para.text.strip()
        if txt:
            content_chunks.append(txt)

    # --- Extract tables ---
    for table_idx, table in enumerate(doc.tables):
        table_md = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_md.append('| ' + ' | '.join(cells) + ' |')
        if table_md:
            table_text = f"\n[Extracted Table {virtual_page_no}.{table_idx+1}]\n" + '\n'.join(table_md)
            content_chunks.append(table_text)

    # --- Extract images with OCR ---
    rels = doc.part.rels
    for rel in rels:
        rel_obj = rels[rel]
        if "image" in rel_obj.target_ref:
            try:
                img_bytes = rel_obj.target_part.blob
                img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
                ocr_text = pytesseract.image_to_string(img)
                if ocr_text.strip():
                    content_chunks.append(f"[OCR from image]\n{ocr_text.strip()}")
            except Exception as e:
                print(f"[OCR failed] {e}", file=sys.stderr)
                content_chunks.append(f"[Image found, but OCR failed: {e}]")

    # --- Optional: Alt text ---
    for shape in getattr(doc, "inline_shapes", []):
        if hasattr(shape, "alt_text") and shape.alt_text:
            content_chunks.append(f"[IMAGE ALT TEXT] {shape.alt_text.strip()}")

    # --- Wrap into a single "virtual page" ---
    if content_chunks:
        full_text = "\n\n".join(content_chunks).strip()
        extracted.append((full_text, virtual_page_no))

    return extracted  # List of (text, page_no)


def extract_text_from_txt(file_path: Path) -> List[Tuple[str, int]]:
    """
    Extracts text from a TXT file, returns as [(text, 1)].
    Deduplicates using MD5 hash (future-proof for multi-page logic).
    
    :param file_path: Path to the TXT file.
    :return: List of (page_text, page_number).
    """
    extracted: List[Tuple[str, int]] = []
    seen_hashes = set()

    try:
        logger.info("Opening TXT %s", file_path)
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read().strip()

            if not text:
                logger.warning("TXT file %s is empty", file_path)
                return []

            hash_val = hashlib.md5(text.encode("utf-8")).hexdigest()
            if hash_val in seen_hashes:
                logger.debug("Duplicate TXT skipped")
                return []

            extracted.append((text, 1))
            seen_hashes.add(hash_val)

    except Exception as e:
        logger.error("Failed extracting TXT %s: %s", file_path, e, exc_info=True)
        return []

    logger.info("Extraction complete—%d page(s)", len(extracted))
    return extracted

def extract_text_from_md(file_path):
    return extract_text_from_txt(file_path)

def extract_text_from_html(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            soup = BeautifulSoup(f, 'html.parser')
            return soup.get_text(separator="\n").strip()
    except Exception as e:
        print(f"Error extracting HTML: {e}", file=sys.stderr)
        return ""

def extract_text_from_csv(file_path):
    try:
        text = ""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            reader = csv.reader(f)
            for row in reader:
                text += ", ".join(row) + "\n"
        return text.strip()
    except Exception as e:
        print(f"Error extracting CSV: {e}", file=sys.stderr)
        return ""

     
# --- Main document loader ---
def load_documents(docs_dir):
    docs_dir = Path(docs_dir)
    extractors = {
        '.pdf': extract_text_from_pdf,
        '.docx': extract_text_from_docx,
        '.txt': extract_text_from_txt,
        '.md': extract_text_from_md,
        '.html': extract_text_from_html,
        '.htm': extract_text_from_html,
        '.csv': extract_text_from_csv,
    }
    documents = []
    for file_path in docs_dir.glob('*'):
        extractor = extractors.get(file_path.suffix.lower())
        if extractor:
            text = extractor(str(file_path))
            if text:
                documents.append(text)
                print(f"\n--- Extracted from {file_path} ---\n{text[:1000]}...\n")
        else:
            print(f"Unsupported file type: {file_path.suffix}", file=sys.stderr)
    return documents

                           

